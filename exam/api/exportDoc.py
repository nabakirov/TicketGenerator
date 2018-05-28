from docx import Document
import os
from exam.configs import FILES_PATH
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX


def getDoc(data):
    document = Document()
    for ticket in data['tickets']:
        if data.get('header'):
            d = document.add_heading(data['header']).add_run()
            d.alignment = WD_ALIGN_PARAGRAPH.CENTER
            d.color = WD_COLOR_INDEX.BLACK
        d2 = document.add_heading('Билет № {}'.format(ticket['t_number'])).add_run()
        d2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for q in ticket['questions']:
            document.add_paragraph('\t{}. {}'.format(q['q_number'], q['text'])).add_run()
        if data.get('footer'):
            d3 = document.add_paragraph(data['footer']).add_run()
            d3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            d3.color = WD_COLOR_INDEX.BLACK
    file_dir = '{}/{}/{}'.format(FILES_PATH, data['user_id'], data['subject_id'])
    os.makedirs(file_dir, exist_ok=True)
    filename = '{}_{}_{}_{}_{}.docx'.format(data['user_id'], data['subject_id'], data['t_count'], data['q_count'], datetime.now().strftime("%d-%m-%Y-%H-%M-%S"))
    file_dir = file_dir + '/' + filename
    document.save(file_dir)
    return filename
